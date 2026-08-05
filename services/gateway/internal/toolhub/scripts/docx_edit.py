
import json, sys, os
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.run import Run
except Exception:
    print(json.dumps({"error":"DOCX adapter requires python-docx"}))
    sys.exit(0)

req = json.load(sys.stdin)
replacements = req.get("replacements") or []
counts = {item["Find"] if "Find" in item else item.get("find"): 0 for item in replacements}

def find_value(item):
    return item["Find"] if "Find" in item else item.get("find", "")

def replace_value(item):
    return item["Replace"] if "Replace" in item else item.get("replace", "")

def run_boundary(run_element, paragraph_element):
    relationship_id = ""
    unsupported = []
    current = run_element.getparent()
    while current is not None and current is not paragraph_element:
        if current.tag == qn("w:hyperlink"):
            relationship_id = current.get(qn("r:id")) or ""
        local_name = current.tag.rsplit("}", 1)[-1]
        if local_name in ("ins", "del", "moveFrom", "moveTo"):
            unsupported.append("tracked_change:%s" % local_name)
        current = current.getparent()
    if run_element.xpath(".//w:fldChar | .//w:instrText"):
        unsupported.append("field")
    if run_element.xpath(".//w:drawing | .//w:pict | .//w:object"):
        unsupported.append("drawing")
    return relationship_id, unsupported

def paragraph_run_spans(paragraph):
    spans = []
    offset = 0
    for run_element in paragraph._p.xpath(".//w:r"):
        run = Run(run_element, paragraph)
        text = run.text or ""
        relationship_id, unsupported = run_boundary(run_element, paragraph._p)
        spans.append({
            "run": run,
            "text": text,
            "start": offset,
            "end": offset + len(text),
            "format": str(run_element.rPr.xml) if run_element.rPr is not None else "",
            "relationship_id": relationship_id,
            "unsupported": unsupported,
        })
        offset += len(text)
    return spans

def paragraph_replacements(text):
    matches = []
    for item in replacements:
        find = find_value(item)
        if not find:
            continue
        cursor = 0
        while True:
            start = text.find(find, cursor)
            if start < 0:
                break
            matches.append({"start": start, "end": start + len(find), "find": find, "replace": replace_value(item)})
            cursor = start + len(find)
    matches.sort(key=lambda match: (match["start"], match["end"], match["find"]))
    for index in range(1, len(matches)):
        if matches[index]["start"] < matches[index - 1]["end"]:
            raise ValueError("DOCX replacement targets overlap in one paragraph")
    return matches

def affected_run_indexes(spans, match):
    text_indexes = [
        index for index, span in enumerate(spans)
        if span["start"] < match["end"] and span["end"] > match["start"]
    ]
    if not text_indexes:
        raise ValueError("DOCX replacement target could not be mapped to text runs")
    return list(range(text_indexes[0], text_indexes[-1] + 1))

def validate_run_span(spans, indexes):
    affected = [spans[index] for index in indexes]
    unsupported = sorted({boundary for span in affected for boundary in span["unsupported"]})
    if unsupported:
        raise ValueError("DOCX replacement crosses unsupported %s boundary" % ", ".join(unsupported))
    text_spans = [span for span in affected if span["text"]]
    if len(text_spans) <= 1:
        return
    formats = {span["format"] for span in text_spans}
    relationships = {span["relationship_id"] for span in text_spans}
    if len(formats) != 1:
        raise ValueError("DOCX replacement crosses mixed run formatting")
    if len(relationships) != 1:
        raise ValueError("DOCX replacement crosses a hyperlink relationship boundary")

def apply_run_replacement(spans, indexes, match):
    first_index = indexes[0]
    last_index = indexes[-1]
    first = spans[first_index]
    last = spans[last_index]
    first_start = match["start"] - first["start"]
    last_end = match["end"] - last["start"]
    if first_index == last_index:
        first["run"].text = first["text"][:first_start] + match["replace"] + first["text"][last_end:]
        return
    first["run"].text = first["text"][:first_start] + match["replace"]
    for index in indexes[1:-1]:
        spans[index]["run"].text = ""
    last["run"].text = last["text"][last_end:]

def replace_in_paragraph(paragraph):
    spans = paragraph_run_spans(paragraph)
    full = "".join(span["text"] for span in spans)
    matches = paragraph_replacements(full)
    if not matches:
        return 0
    prepared = []
    for match in matches:
        indexes = affected_run_indexes(spans, match)
        validate_run_span(spans, indexes)
        prepared.append((match, indexes))
    for match, indexes in reversed(prepared):
        apply_run_replacement(spans, indexes, match)
        counts[match["find"]] = counts.get(match["find"], 0) + 1
    return len(prepared)

try:
    doc = Document(req["path"])
    total = 0
    for paragraph in doc.paragraphs:
        total += replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total += replace_in_paragraph(paragraph)
    missing = [find for find, count in counts.items() if count == 0]
    if missing:
        print(json.dumps({"error":"find text was not matched: " + ", ".join(repr(x) for x in missing)}))
        sys.exit(0)
    os.makedirs(os.path.dirname(req["output_path"]), exist_ok=True)
    doc.save(req["output_path"])
    print(json.dumps({"replacements":total,"bytes":os.path.getsize(req["output_path"]),"details":[{"find":k,"count":v} for k,v in counts.items()]}))
except Exception as e:
    print(json.dumps({"error":str(e)}))
