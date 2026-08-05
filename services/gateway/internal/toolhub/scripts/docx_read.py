import base64, json, sys
try:
    import docx
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
except Exception:
    print(json.dumps({"error":"DOCX reader requires python-docx"}))
    sys.exit(0)

req = json.load(sys.stdin)
max_bytes = int(req.get("max_bytes") or 20000)

def trim(text):
    return " ".join(str(text or "").split())

def enum_value(value):
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    return str(value)

def length_value(value):
    return int(value) if value is not None else 0

def point_value(value):
    if value is None:
        return None
    return round(float(value.pt), 4)

def color_value(font):
    try:
        if font.color.rgb is not None:
            return str(font.color.rgb)
        if font.color.theme_color is not None:
            return "theme:%s" % font.color.theme_color
    except Exception:
        pass
    return None

def style_font_value(style, field):
    visited = set()
    while style is not None and id(style) not in visited:
        visited.add(id(style))
        try:
            value = getattr(style.font, field)
            if value is not None:
                return value
            style = style.base_style
        except Exception:
            return None
    return None

def effective_run_value(run, paragraph, field):
    try:
        value = getattr(run.font, field)
        if value is not None:
            return value
    except Exception:
        pass
    for style in (getattr(run, "style", None), getattr(paragraph, "style", None)):
        value = style_font_value(style, field)
        if value is not None:
            return value
    return None

def ancestor_metadata(element, paragraph_element):
    relationship_id = ""
    tracked_change = ""
    current = element.getparent()
    while current is not None and current is not paragraph_element:
        if current.tag == qn("w:hyperlink"):
            relationship_id = current.get(qn("r:id")) or ""
        local_name = current.tag.rsplit("}", 1)[-1]
        if local_name in ("ins", "del", "moveFrom", "moveTo"):
            tracked_change = local_name
        current = current.getparent()
    return relationship_id, tracked_change

def run_fields(run, paragraph, run_element, start, end):
    relationship_id, tracked_change = ancestor_metadata(run_element, paragraph._p)
    boundaries = []
    if relationship_id:
        boundaries.append("hyperlink")
    if tracked_change:
        boundaries.append("tracked_change:%s" % tracked_change)
    if run_element.xpath(".//w:fldChar | .//w:instrText"):
        boundaries.append("field")
    if run_element.xpath(".//w:drawing | .//w:pict | .//w:object"):
        boundaries.append("drawing")
    effective_bold = effective_run_value(run, paragraph, "bold")
    effective_size = effective_run_value(run, paragraph, "size")
    return {
        "text": run.text or "",
        "start": start,
        "end": end,
        "bold": run.bold,
        "italic": run.italic,
        "underline": enum_value(run.underline),
        "font_name": run.font.name,
        "font_size_pt": point_value(run.font.size),
        "font_color": color_value(run.font),
        "effective_bold": bool(effective_bold) if effective_bold is not None else None,
        "effective_font_size_pt": point_value(effective_size),
        "relationship_id": relationship_id,
        "boundaries": boundaries,
    }

def paragraph_runs(paragraph):
    runs = []
    offset = 0
    for index, run_element in enumerate(paragraph._p.xpath(".//w:r"), start=1):
        run = Run(run_element, paragraph)
        text = run.text or ""
        item = run_fields(run, paragraph, run_element, offset, offset + len(text))
        item["index"] = index
        runs.append(item)
        offset += len(text)
    return runs

def paragraph_boundaries(paragraph):
    boundaries = []
    checks = (
        ("field", ".//w:fldSimple | .//w:fldChar | .//w:instrText"),
        ("drawing", ".//w:drawing | .//w:pict | .//w:object"),
        ("tracked_change", ".//w:ins | .//w:del | .//w:moveFrom | .//w:moveTo"),
        ("content_control", ".//w:sdt"),
        ("text_box", ".//w:txbxContent"),
    )
    for name, expression in checks:
        if paragraph._p.xpath(expression):
            boundaries.append(name)
    return boundaries

def paragraph_fields(paragraph):
    style = paragraph.style.name if paragraph.style is not None else ""
    outline_level = ""
    try:
        outline_level = enum_value(paragraph.paragraph_format.outline_level) or ""
    except Exception:
        pass
    num_id = ""
    try:
        num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
        if num_pr is not None and num_pr.numId is not None:
            num_id = str(num_pr.numId.val)
    except Exception:
        pass
    paragraph_format = paragraph.paragraph_format
    formatting = {
        "alignment": enum_value(paragraph.alignment),
        "left_indent": length_value(paragraph_format.left_indent),
        "right_indent": length_value(paragraph_format.right_indent),
        "first_line_indent": length_value(paragraph_format.first_line_indent),
        "space_before": length_value(paragraph_format.space_before),
        "space_after": length_value(paragraph_format.space_after),
        "line_spacing": enum_value(paragraph_format.line_spacing),
        "keep_together": paragraph_format.keep_together,
        "keep_with_next": paragraph_format.keep_with_next,
        "page_break_before": paragraph_format.page_break_before,
        "widow_control": paragraph_format.widow_control,
    }
    return style, outline_level, num_id, formatting

def paragraph_record(paragraph, index, part_kind, location):
    runs = paragraph_runs(paragraph)
    raw_text = "".join(run["text"] for run in runs)
    if raw_text == "":
        raw_text = paragraph.text or ""
    style, outline_level, num_id, formatting = paragraph_fields(paragraph)
    return {
        "index": index,
        "text": trim(raw_text),
        "raw_text": raw_text,
        "style": style,
        "outline_level": outline_level,
        "list_id": num_id,
        "part_kind": part_kind,
        "format": formatting,
        "runs": runs,
        "unsupported_boundaries": paragraph_boundaries(paragraph),
        "location": location,
    }

def paragraph_hyperlinks(paragraph, location, part, hyperlinks):
    for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
        rel_id = hyperlink.get(qn("r:id")) or ""
        target = ""
        if rel_id and rel_id in part.rels:
            target = str(part.rels[rel_id].target_ref)
        hyperlinks.append({
            "kind": "hyperlink",
            "text": trim("".join((Run(run_element, paragraph).text or "") for run_element in hyperlink.xpath(".//w:r"))),
            "target": target,
            "location": dict(location),
            "source": {"parser": "python_docx", "relationship_id": rel_id},
        })

def paragraph_image_rel_ids(paragraph):
    return [value for value in paragraph._p.xpath(".//a:blip/@r:embed") if value]

def part_root(part):
    root = getattr(part, "element", None) or getattr(part, "_element", None)
    if root is not None:
        return root
    try:
        return parse_xml(part.blob)
    except Exception:
        return None

def root_has_any(root, names):
    if root is None:
        return False
    for element in root.iter():
        if element.tag.rsplit("}", 1)[-1] in names:
            return True
    return False

extracted_bytes = 0

def append_line(lines, text):
    global extracted_bytes
    extracted_bytes += len(text.encode("utf-8")) + (1 if lines else 0)
    if extracted_bytes > max_bytes:
        print(json.dumps({"content":"", "truncated":True, "extracted_bytes":extracted_bytes}))
        sys.exit(0)
    lines.append(text)

try:
    document = docx.Document(req["path"])
    paragraphs = []
    blocks = []
    tables = []
    lines = []
    block_index = 0
    images = []
    resources = []
    resource_keys = set()
    registered_locations = set()
    hyperlinks = []

    def register_image(part, rel_id, parent_path, location):
        if not rel_id or rel_id not in part.related_parts:
            return
        related = part.related_parts[rel_id]
        content_type = str(getattr(related, "content_type", "application/octet-stream"))
        if not content_type.startswith("image/"):
            return
        part_name = str(getattr(related, "partname", ""))
        owner_part = str(getattr(part, "partname", ""))
        resource_key = "%s:%s:%s" % (owner_part, rel_id, part_name)
        location_key = "%s:%s" % (resource_key, location.get("path", ""))
        if location_key in registered_locations:
            return
        registered_locations.add(location_key)
        blob = bytes(related.blob)
        images.append({
            "kind": "image",
            "resource_key": resource_key,
            "parent_path": parent_path,
            "location": location,
            "source": {
                "parser": "python_docx",
                "relationship_id": rel_id,
                "part_name": part_name,
                "owner_part": owner_part,
            },
            "content_type": content_type,
        })
        if resource_key not in resource_keys:
            resource_keys.add(resource_key)
            resources.append({
                "key": resource_key,
                "kind": "image",
                "content_type": content_type,
                "data_base64": base64.b64encode(blob).decode("ascii"),
            })

    for index, paragraph in enumerate(document.paragraphs, start=1):
        path = "document.p[%d]" % index
        location = {
            "part": "document",
            "part_kind": "body",
            "block_type": "paragraph",
            "paragraph_index": index,
            "table_index": 0,
            "row_index": 0,
            "cell_index": 0,
            "cell_paragraph_index": 0,
            "path": path,
        }
        for image_index, rel_id in enumerate(paragraph_image_rel_ids(paragraph), start=1):
            register_image(document.part, rel_id, path, dict(location, block_type="image", image_index=image_index, path=path+".image[%d]" % image_index))
        paragraph_hyperlinks(paragraph, location, document.part, hyperlinks)
        item = paragraph_record(paragraph, index, "body", location)
        paragraphs.append(item)
        if not item["text"]:
            continue
        block_index += 1
        location["block_index"] = block_index
        blocks.append({"text": item["text"], "style": item["style"], "level": item["outline_level"], "location": location})
        append_line(lines, item["text"])

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row_index, row in enumerate(table.rows, start=1):
            row_values = []
            for cell_index, cell in enumerate(row.cells, start=1):
                cell_texts = []
                for cell_paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                    path = "document.table[%d].row[%d].cell[%d].p[%d]" % (table_index, row_index, cell_index, cell_paragraph_index)
                    location = {
                        "part": "document",
                        "part_kind": "table_cell",
                        "block_type": "table_cell",
                        "paragraph_index": 0,
                        "table_index": table_index,
                        "row_index": row_index,
                        "cell_index": cell_index,
                        "cell_paragraph_index": cell_paragraph_index,
                        "path": path,
                    }
                    for image_index, rel_id in enumerate(paragraph_image_rel_ids(paragraph), start=1):
                        register_image(document.part, rel_id, path, dict(location, block_type="image", image_index=image_index, path=path+".image[%d]" % image_index))
                    paragraph_hyperlinks(paragraph, location, document.part, hyperlinks)
                    item = paragraph_record(paragraph, cell_paragraph_index, "table_cell", location)
                    text = item["text"]
                    if not text:
                        continue
                    cell_texts.append(text)
                    block_index += 1
                    location["block_index"] = block_index
                    blocks.append({"text": text, "style": item["style"], "level": item["outline_level"], "runs": item["runs"], "location": location})
                    append_line(lines, text)
                row_values.append("\n".join(cell_texts))
            rows.append({"index": row_index, "cells": row_values})
        tables.append({"index": table_index, "rows": rows})

    section_layouts = []
    story_references = {}
    current_story_parts = {}
    for section_index, section in enumerate(document.sections, start=1):
        section_path = "document.section[%d]" % section_index
        section_layouts.append({
            "index": section_index,
            "path": section_path,
            "orientation": enum_value(section.orientation),
            "start_type": enum_value(section.start_type),
            "page_width": length_value(section.page_width),
            "page_height": length_value(section.page_height),
            "top_margin": length_value(section.top_margin),
            "right_margin": length_value(section.right_margin),
            "bottom_margin": length_value(section.bottom_margin),
            "left_margin": length_value(section.left_margin),
            "header_linked_to_previous": bool(section.header.is_linked_to_previous),
            "footer_linked_to_previous": bool(section.footer.is_linked_to_previous),
        })
        explicit = {}
        for reference in section._sectPr.xpath("./w:headerReference | ./w:footerReference"):
            kind = "header" if reference.tag == qn("w:headerReference") else "footer"
            variant = reference.get(qn("w:type")) or "default"
            rel_id = reference.get(qn("r:id")) or ""
            part = document.part.related_parts.get(rel_id)
            if part is not None:
                explicit[(kind, variant)] = part
        for kind in ("header", "footer"):
            for variant in ("default", "first", "even"):
                key = (kind, variant)
                if key in explicit:
                    current_story_parts[key] = explicit[key]
                part = current_story_parts.get(key)
                if part is None:
                    continue
                part_name = str(part.partname)
                story_references.setdefault(part_name, {"part": part, "kind": kind, "references": []})
                story_references[part_name]["references"].append({"section_index": section_index, "variant": variant})

    story_parts = []
    for part_name in sorted(story_references):
        story = story_references[part_name]
        part = story["part"]
        kind = story["kind"]
        story_paragraphs = []
        story_blocks = []
        story_tables = []
        root = part_root(part)
        paragraph_elements = list(getattr(root, "p_lst", []))
        table_elements = list(getattr(root, "tbl_lst", []))
        for paragraph_index, paragraph_element in enumerate(paragraph_elements, start=1):
            paragraph = Paragraph(paragraph_element, part)
            path = "document.%s[%s].p[%d]" % (kind, part_name, paragraph_index)
            location = {
                "part": "document", "part_kind": kind, "block_type": "story_paragraph",
                "paragraph_index": 0, "story_paragraph_index": paragraph_index,
                "story_part": part_name, "path": path,
            }
            for image_index, rel_id in enumerate(paragraph_image_rel_ids(paragraph), start=1):
                register_image(part, rel_id, path, dict(location, block_type="image", image_index=image_index, path=path+".image[%d]" % image_index))
            paragraph_hyperlinks(paragraph, location, part, hyperlinks)
            item = paragraph_record(paragraph, paragraph_index, kind, location)
            story_paragraphs.append(item)
            if item["text"]:
                story_blocks.append({"text": item["text"], "style": item["style"], "level": item["outline_level"], "location": location})
        for table_index, table_element in enumerate(table_elements, start=1):
            table = Table(table_element, part)
            rows = []
            for row_index, row in enumerate(table.rows, start=1):
                row_values = []
                for cell_index, cell in enumerate(row.cells, start=1):
                    cell_values = []
                    for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                        path = "document.%s[%s].table[%d].row[%d].cell[%d].p[%d]" % (kind, part_name, table_index, row_index, cell_index, paragraph_index)
                        location = {
                            "part": "document", "part_kind": kind, "block_type": "story_table_cell",
                            "paragraph_index": 0, "story_paragraph_index": paragraph_index,
                            "story_part": part_name, "table_index": table_index, "row_index": row_index,
                            "cell_index": cell_index, "path": path,
                        }
                        paragraph_hyperlinks(paragraph, location, part, hyperlinks)
                        item = paragraph_record(paragraph, paragraph_index, kind, location)
                        if item["text"]:
                            cell_values.append(item["text"])
                            story_blocks.append({"text": item["text"], "style": item["style"], "level": item["outline_level"], "location": location})
                    row_values.append("\n".join(cell_values))
                rows.append({"index": row_index, "cells": row_values})
            story_tables.append({"index": table_index, "rows": rows})
        story_parts.append({
            "kind": kind,
            "part_name": part_name,
            "references": story["references"],
            "paragraphs": story_paragraphs,
            "tables": story_tables,
            "blocks": story_blocks,
        })

    for rel_id, relation in document.part.rels.items():
        if str(relation.reltype).endswith("/image"):
            related = document.part.related_parts.get(rel_id)
            resource_key = "%s:%s:%s" % (str(document.part.partname), rel_id, str(getattr(related, "partname", "")))
            if resource_key not in resource_keys:
                register_image(document.part, rel_id, "document", {
                    "part": "document", "part_kind": "body", "block_type": "image",
                    "path": "document.relationship[%s]" % rel_id,
                })

    comments = []
    try:
        for comment in document.comments:
            comments.append({
                "kind": "comment",
                "text": trim(comment.text),
                "author": str(comment.author or ""),
                "initials": str(comment.initials or ""),
                "location": {"path": "document.comment[%s]" % comment.comment_id},
                "source": {"parser": "python_docx", "comment_id": int(comment.comment_id)},
            })
    except Exception:
        pass

    omissions = []
    part_names = []
    story_names = {item["part_name"] for item in story_parts}
    for part in document.part.package.parts:
        part_name = str(part.partname)
        part_names.append(part_name)
        root = part_root(part)
        reasons = []
        lower_name = part_name.lower()
        if lower_name.endswith("/footnotes.xml"):
            reasons.append("footnotes")
        if lower_name.endswith("/endnotes.xml"):
            reasons.append("endnotes")
        if root_has_any(root, {"txbxContent"}):
            reasons.append("text_boxes")
        if root_has_any(root, {"ins", "del", "moveFrom", "moveTo"}):
            reasons.append("tracked_changes")
        if root_has_any(root, {"altChunk"}):
            reasons.append("alt_chunk")
        if root_has_any(root, {"sdt"}):
            reasons.append("content_controls")
        if root_has_any(root, {"tbl"}):
            for table_element in root.iter(qn("w:tbl")):
                parent = table_element.getparent()
                while parent is not None and parent is not root:
                    if parent.tag == qn("w:tc"):
                        reasons.append("nested_tables")
                        parent = None
                        break
                    parent = parent.getparent()
                if "nested_tables" in reasons:
                    break
        known_text_part = lower_name.endswith("/document.xml") or part_name in story_names or lower_name.endswith("/comments.xml")
        if root_has_any(root, {"t"}) and lower_name.startswith("/word/") and not known_text_part and not reasons:
            reasons.append("unrecognized_text_part")
        for reason in sorted(set(reasons)):
            omissions.append({"part_name": part_name, "reason": reason})

    omission_reasons = {item["reason"] for item in omissions}
    omissions_by_part = {}
    for omission in omissions:
        omissions_by_part.setdefault(omission["part_name"], []).append(omission["reason"])
    for story in story_parts:
        story["coverage"] = "partial" if story["part_name"] in omissions_by_part else "complete"
    content_status = "partial" if omissions else "complete"
    content_scopes = {
        "body": "partial" if str(document.part.partname) in omissions_by_part else "complete",
        "tables": "partial" if "nested_tables" in omission_reasons else "complete",
        "headers": "partial" if any(item["kind"] == "header" and item["coverage"] == "partial" for item in story_parts) else "complete",
        "footers": "partial" if any(item["kind"] == "footer" and item["coverage"] == "partial" for item in story_parts) else "complete",
        "footnotes": "unsupported" if "footnotes" in omission_reasons else "absent",
        "endnotes": "unsupported" if "endnotes" in omission_reasons else "absent",
        "text_boxes": "unsupported" if "text_boxes" in omission_reasons else "absent",
        "tracked_changes": "partial" if "tracked_changes" in omission_reasons else "absent",
    }

    content = "\n".join(lines).strip()
    raw = content.encode("utf-8")
    truncated = len(raw) > max_bytes
    if truncated:
        content = raw[:max_bytes].decode("utf-8", errors="ignore")
    print(json.dumps({
        "content": content,
        "truncated": truncated,
        "extracted_bytes": extracted_bytes,
        "resources": resources,
        "document": {
            "schema_version": "document_read_v1",
            "format": "docx",
            "source": "python_docx",
            "blocks": blocks,
            "paragraphs": paragraphs,
            "tables": tables,
            "enrichment": {
                "schema_version": "document_enrichment_v1",
                "assets": {"images": images, "charts": [], "embedded_objects": []},
                "annotations": {"comments": comments, "notes": [], "hyperlinks": hyperlinks},
                "layout": {"sections": section_layouts, "page_settings": section_layouts, "slide_layouts": [], "merged_ranges": []},
                "extensions": {
                    "status": "partial" if omissions else "complete",
                    "parts": sorted(part_names),
                    "story_parts": story_parts,
                    "unparsed_parts": sorted({item["part_name"] for item in omissions}),
                    "content_omissions": omissions,
                },
                "coverage": {
                    "content": content_status,
                    "content_scopes": content_scopes,
                    "assets": "complete",
                    "annotations": "partial",
                    "layout": "partial",
                    "extensions": "partial" if omissions else "complete",
                },
            },
            "stats": {
                "blocks": len(blocks),
                "paragraphs": len(paragraphs),
                "tables": len(tables),
                "story_parts": len(story_parts),
                "images": len(images),
                "comments": len(comments),
                "hyperlinks": len(hyperlinks),
                "complete": not truncated,
            }
        }
    }, ensure_ascii=False))
except Exception as e:
    print(json.dumps({"error":str(e)}))
