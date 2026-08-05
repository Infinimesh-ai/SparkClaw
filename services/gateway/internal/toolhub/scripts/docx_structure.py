
import hashlib, json, sys, os
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.text.run import Run
except Exception:
    print(json.dumps({"error":"DOCX structure adapter requires python-docx"}))
    sys.exit(0)

req = json.load(sys.stdin)
op = req.get("operation")

def trim(text):
    return " ".join(str(text or "").split())

def source_hash(text):
    return "sha1:" + hashlib.sha1(trim(text).encode("utf-8")).hexdigest()

def requested_location():
    loc = req.get("location")
    if loc in (None, ""):
        return None
    if not isinstance(loc, dict):
        raise ValueError("location must be an object")
    part = str(loc.get("part") or "document")
    block_type = str(loc.get("block_type") or "")
    if part != "document":
        raise ValueError("only document part locations are currently editable")
    if block_type != "paragraph":
        raise ValueError("only top-level paragraph locations are currently editable")
    idx = int(loc.get("paragraph_index") or 0)
    if idx <= 0:
        raise ValueError("location.paragraph_index must be a positive 1-based integer")
    return loc

def paragraph_index():
    loc = requested_location()
    if loc is not None:
        return int(loc.get("paragraph_index") or 0)
    idx = int(req.get("paragraph_index") or 0)
    if idx <= 0:
        raise ValueError("paragraph_index or location must identify a positive 1-based paragraph")
    return idx

def paragraph_at(doc, idx):
    if idx < 1 or idx > len(doc.paragraphs):
        raise ValueError("paragraph_index out of range: %s" % idx)
    return doc.paragraphs[idx - 1]

def preflight_paragraph(paragraph):
    before = trim(paragraph.text)
    expected = trim(req.get("old_text") or "")
    if expected and before != expected:
        raise ValueError("old_text mismatch at target paragraph")
    expected_hash = str(req.get("source_hash") or "").strip()
    actual_hash = source_hash(before)
    if expected_hash and actual_hash != expected_hash:
        raise ValueError("source_hash mismatch at target paragraph")
    return before, actual_hash

def replacement_run_boundary(run_element, paragraph_element):
    relationship_id = ""
    unsupported = []
    current = run_element.getparent()
    while current is not None and current is not paragraph_element:
        if current.tag == qn("w:hyperlink"):
            relationship_id = current.get(qn("r:id")) or ""
        local_name = current.tag.rsplit("}", 1)[-1]
        if local_name in ("ins", "del", "moveFrom", "moveTo"):
            unsupported.append("tracked_change:%s" % local_name)
        current = current.getparent()
    if run_element.xpath(".//w:fldChar | .//w:instrText"):
        unsupported.append("field")
    if run_element.xpath(".//w:drawing | .//w:pict | .//w:object"):
        unsupported.append("drawing")
    return relationship_id, unsupported

def homogeneous_replacement_runs(paragraph):
    runs = []
    for run_element in paragraph._p.xpath(".//w:r"):
        run = Run(run_element, paragraph)
        relationship_id, unsupported = replacement_run_boundary(run_element, paragraph._p)
        if unsupported:
            raise ValueError("DOCX paragraph replacement crosses unsupported %s boundary" % ", ".join(sorted(set(unsupported))))
        runs.append({
            "run": run,
            "text": run.text or "",
            "format": str(run_element.rPr.xml) if run_element.rPr is not None else "",
            "relationship_id": relationship_id,
        })
    text_runs = [item for item in runs if item["text"]]
    if len({item["format"] for item in text_runs}) > 1:
        raise ValueError("DOCX paragraph replacement requires homogeneous run formatting")
    if len({item["relationship_id"] for item in text_runs}) > 1:
        raise ValueError("DOCX paragraph replacement crosses a hyperlink relationship boundary")
    return runs, text_runs

def set_paragraph_text(paragraph, text):
    runs, text_runs = homogeneous_replacement_runs(paragraph)
    if not runs:
        paragraph.add_run(text)
        return
    target = text_runs[0] if text_runs else runs[0]
    for item in runs:
        item["run"].text = ""
    target["run"].text = text

def insert_paragraph(doc, position, idx, text):
    position = (position or "").strip().lower()
    if position == "start":
        if doc.paragraphs:
            return doc.paragraphs[0].insert_paragraph_before(text)
        return doc.add_paragraph(text)
    if position == "end":
        return doc.add_paragraph(text)
    if position in ("before", "after"):
        paragraph = paragraph_at(doc, idx)
        if position == "before":
            return paragraph.insert_paragraph_before(text)
        inserted = paragraph.insert_paragraph_before(text)
        paragraph._p.addnext(inserted._p)
        return inserted
    raise ValueError("position must be one of start, end, before, after")

def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None

def apply_style(paragraph, style_req):
    if not isinstance(style_req, dict):
        raise ValueError("style must be an object")
    applied = {}
    builtin_style = str(style_req.get("builtin_style") or "").strip()
    if builtin_style:
        paragraph.style = builtin_style
        applied["builtin_style"] = builtin_style
    bold = style_req.get("bold")
    font_size = style_req.get("font_size_pt")
    if bold is not None or font_size is not None:
        if not paragraph.runs:
            paragraph.add_run("")
        for run in paragraph.runs:
            if bold is not None:
                run.bold = bool(bold)
            if font_size is not None:
                size = int(font_size)
                if size <= 0 or size > 200:
                    raise ValueError("font_size_pt must be between 1 and 200")
                run.font.size = Pt(size)
        if bold is not None:
            applied["bold"] = bool(bold)
        if font_size is not None:
            applied["font_size_pt"] = int(font_size)
    if not applied:
        raise ValueError("style must contain builtin_style, bold, or font_size_pt")
    return applied

try:
    doc = Document(req["path"])
    text = str(req.get("text") or "")
    result = {
        "status": "docx_version_written",
        "operation": op,
        "path": req["path"],
        "output_path": req["output_path"]
    }
    loc = requested_location()
    if loc is not None:
        result["location"] = loc
    if op == "replace_paragraph":
        idx = paragraph_index()
        paragraph = paragraph_at(doc, idx)
        before, before_hash = preflight_paragraph(paragraph)
        set_paragraph_text(paragraph, text)
        result["paragraph_index"] = idx
        result["before"] = before
        result["source_hash"] = before_hash
        result["text"] = text
    elif op == "insert_paragraph":
        position = str(req.get("position") or "").strip().lower()
        idx = paragraph_index() if loc is not None else int(req.get("paragraph_index") or 0)
        if position in ("before", "after") and idx <= 0:
            raise ValueError("paragraph_index or location is required for before/after insertion")
        if position in ("before", "after"):
            anchor = paragraph_at(doc, idx)
            before, before_hash = preflight_paragraph(anchor)
            result["anchor_before"] = before
            result["source_hash"] = before_hash
        insert_paragraph(doc, position, idx, text)
        result["position"] = position
        result["text"] = text
        if position == "start":
            result["paragraph_index"] = 1
        elif position == "end":
            result["paragraph_index"] = len(doc.paragraphs)
        elif position == "before":
            result["paragraph_index"] = idx
        else:
            result["paragraph_index"] = idx + 1
    elif op == "delete_paragraph":
        idx = paragraph_index()
        paragraph = paragraph_at(doc, idx)
        before, before_hash = preflight_paragraph(paragraph)
        result["paragraph_index"] = idx
        result["text"] = before
        result["source_hash"] = before_hash
        delete_paragraph(paragraph)
    elif op == "set_text_style":
        idx = paragraph_index()
        paragraph = paragraph_at(doc, idx)
        before, before_hash = preflight_paragraph(paragraph)
        applied = apply_style(paragraph, req.get("style"))
        result["paragraph_index"] = idx
        result["before"] = before
        result["source_hash"] = before_hash
        result["style"] = applied
    else:
        raise ValueError("unsupported docx operation: %s" % op)
    os.makedirs(os.path.dirname(req["output_path"]), exist_ok=True)
    doc.save(req["output_path"])
    result["bytes"] = os.path.getsize(req["output_path"])
    print(json.dumps(result, ensure_ascii=False))
except Exception as e:
    print(json.dumps({"error":str(e)}, ensure_ascii=False))
